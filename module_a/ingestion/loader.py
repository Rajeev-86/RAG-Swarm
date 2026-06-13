"""
module_a/ingestion/loader.py
─────────────────────────────
Loads PDF, DOCX, TXT, and MD files from a data room directory.
Auto-infers M&A domain (legal / financial / hr / cybersecurity) from
filename and a content snippet — no manual tagging required.
"""

import os
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional


# ── Data Model ───────────────────────────────────────────────────────────────

@dataclass
class RawDocument:
    content:   str
    source:    str
    domain:    str         # legal | financial | hr | cybersecurity | unknown
    file_type: str
    metadata:  dict = field(default_factory=dict)


# ── Domain Keyword Heuristics ─────────────────────────────────────────────────

DOMAIN_KEYWORDS: dict[str, list[str]] = {
    "legal": [
        "contract", "agreement", "clause", "regulatory", "compliance",
        "litigation", "ip", "patent", "indemnif", "arbitration",
        "jurisdiction", "liability", "escrow", "covenant", "warranty",
    ],
    "financial": [
        "revenue", "ebitda", "balance", "cashflow", "audit", "tax",
        "financial", "p&l", "income", "expenses", "margin", "valuation",
        "debt", "equity", "working capital", "capex",
    ],
    "hr": [
        "employee", "headcount", "compensation", "benefit", "payroll",
        "talent", "org", "workforce", "attrition", "retention",
        "bonus", "salary", "vesting", "option", "hire",
    ],
    "cybersecurity": [
        "security", "vulnerability", "infosec", "penetration", "cve",
        "breach", "firewall", "soc2", "encryption", "access control",
        "incident", "patch", "gdpr", "iso27001", "zero day",
    ],
}


def infer_domain(filename: str, content_snippet: str = "") -> str:
    """
    Infer the M&A domain from the filename and optionally a content snippet.
    Uses a simple keyword-count heuristic — good enough for routing;
    the cross-encoder reranker provides precise filtering later.
    """
    combined = (filename + " " + content_snippet[:600]).lower()
    scores = {
        domain: sum(kw in combined for kw in keywords)
        for domain, keywords in DOMAIN_KEYWORDS.items()
    }
    best_domain = max(scores, key=scores.get)
    return best_domain if scores[best_domain] > 0 else "unknown"


# ── Format-Specific Loaders ───────────────────────────────────────────────────

def _load_pdf(path: str) -> str:
    from pypdf import PdfReader
    reader = PdfReader(path)
    return "\n".join(
        page.extract_text() or ""
        for page in reader.pages
    )


def _load_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _load_text(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


_LOADERS = {
    ".pdf":  _load_pdf,
    ".docx": _load_docx,
    ".txt":  _load_text,
    ".md":   _load_text,
}


# ── Public API ────────────────────────────────────────────────────────────────

def load_document(path: str) -> Optional[RawDocument]:
    """Load a single file and return a RawDocument, or None if unsupported / failed."""
    p = Path(path)
    ext = p.suffix.lower()
    loader = _LOADERS.get(ext)
    if not loader:
        return None
    try:
        content = loader(path)
        if not content.strip():
            return None
        domain = infer_domain(p.name, content)
        return RawDocument(
            content=content,
            source=str(p.resolve()),
            domain=domain,
            file_type=ext.lstrip("."),
            metadata={
                "filename":    p.name,
                "domain":      domain,
                "size_chars":  len(content),
            },
        )
    except Exception as exc:
        print(f"[Loader] ⚠  Failed to load {path}: {exc}")
        return None


def load_data_room(data_room_dir: str) -> list[RawDocument]:
    """
    Recursively walk a data room directory and load every supported document.
    Returns a list of RawDocument objects, skipping unsupported / empty files.
    """
    docs: list[RawDocument] = []
    skipped = 0
    for root, _, files in os.walk(data_room_dir):
        for filename in sorted(files):
            full_path = os.path.join(root, filename)
            doc = load_document(full_path)
            if doc:
                docs.append(doc)
            else:
                skipped += 1

    print(
        f"[Loader] ✓ Loaded {len(docs)} documents "
        f"({skipped} skipped) from '{data_room_dir}'"
    )
    return docs